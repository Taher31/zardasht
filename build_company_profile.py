from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path('Asperlos_Gostar_Fardad_Company_Profile_EN_FA.docx')

NAVY = '17324D'
GOLD = 'B88A2B'
LIGHT = 'F3F6F7'
GRAY = '5B6870'

def set_font(run, name='Arial', size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:cs'), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_page_number(paragraph):
    run = paragraph.add_run('Page ')
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

def set_cell_text(cell, text, bold=False, rtl_text=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if rtl_text:
        rtl(p)
    r = p.add_run(text)
    set_font(r, size=10, color=NAVY, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_bullet(doc, text, rtl_text=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if rtl_text:
        rtl(p)
    r = p.add_run(text)
    set_font(r, size=10.5, color='222222')
    return p

def add_heading(doc, text, level=1, rtl_text=False):
    p = doc.add_paragraph(style=f'Heading {level}')
    if rtl_text:
        rtl(p)
    r = p.add_run(text)
    set_font(r, size={1:16,2:13,3:12}[level], color=NAVY, bold=True)
    return p

def add_body(doc, text, rtl_text=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    if rtl_text:
        rtl(p)
    r = p.add_run(text)
    set_font(r, size=10.5, color='222222')
    return p

def add_label_value(doc, rows, rtl_text=False):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.7)
        cells[1].width = Inches(4.8)
        shade(cells[0], LIGHT)
        set_cell_text(cells[0], label, bold=True, rtl_text=rtl_text)
        set_cell_text(cells[1], value, rtl_text=rtl_text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10.5)
styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
styles['Normal']._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
for name, size, before, after in [('Heading 1',16,16,7),('Heading 2',13,11,5),('Heading 3',12,8,4)]:
    s=styles[name]
    s.font.name='Arial'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(NAVY)
    s._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); s._element.rPr.rFonts.set(qn('w:cs'),'Arial')
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

# Footer
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(fp)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(70)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('COMPANY PROFILE'); set_font(r,size=11,color=GOLD,bold=True)
p.paragraph_format.space_after=Pt(10)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Asperlos Gostar Fardad Trading'); set_font(r,size=28,color=NAVY,bold=True)
p.paragraph_format.space_after=Pt(8)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Pulses & Legumes Exporters'); set_font(r,size=15,color=GRAY)
p.paragraph_format.space_after=Pt(34)
p=doc.add_paragraph(); rtl(p)
r=p.add_run('بازرگانی اسپرلوس گستر فرداد'); set_font(r,size=24,color=NAVY,bold=True)
p.paragraph_format.space_after=Pt(6)
p=doc.add_paragraph(); rtl(p)
r=p.add_run('صادرکننده حبوبات و محصولات بقولات'); set_font(r,size=14,color=GRAY)
p.paragraph_format.space_after=Pt(56)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Iran | B2B Trade | Export-Ready Supply'); set_font(r,size=10.5,color=GOLD,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Bilingual corporate profile - English / فارسی'); set_font(r,size=10,color=GRAY)

doc.add_page_break()

# English profile
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=p.add_run('ENGLISH'); set_font(r,size=10,color=GOLD,bold=True)
add_heading(doc,'Company Overview')
add_body(doc,'Asperlos Gostar Fardad Trading is an Iran-based B2B trading company focused on the sourcing, grading, packaging, and export of pulses and legumes. We connect Iranian growers and packhouses with importers, wholesalers, retail packers, and food manufacturers in the Middle East, Asia, and Africa.')
add_body(doc,'Our purpose is to make Iranian pulses easier to import by combining consistent product quality with transparent documentation, practical logistics coordination, and dependable communication from inquiry to delivery.')
add_label_value(doc,[
    ('Company', 'Asperlos Gostar Fardad Trading'),
    ('Industry', 'Agricultural commodities and food ingredients'),
    ('Core business', 'Sourcing and export of pulses and legumes'),
    ('Base', 'Iran'),
    ('Markets served', 'Middle East, Asia, and Africa'),
    ('Customer types', 'Importers, traders, wholesalers, retail packers, canneries, food manufacturers, and foodservice distributors'),
])

add_heading(doc,'Mission, Vision and Values')
add_heading(doc,'Mission',2)
add_body(doc,'To make Iranian pulses easy to import through consistent quality, transparent paperwork, and dependable schedules.')
add_heading(doc,'Vision',2)
add_body(doc,'To be a trusted long-term trade partner for buyers who need reliable Iranian pulse supply, clear specifications, and export execution that supports repeat business.')
add_heading(doc,'Values',2)
for item in ['Quality first, for every container.', 'On-time shipment with proactive milestone updates.', 'Transparent documentation and compliance support.', 'Long-term partnerships over one-off transactions.', 'Responsive communication across time zones.']:
    add_bullet(doc,item)

add_heading(doc,'Product Portfolio')
add_body(doc,'We supply export-ready legumes tailored to agreed specifications, packaging requirements, and destination-market needs.')
add_heading(doc,'Chickpeas',2)
add_body(doc,'Kabuli chickpeas are supplied with machine cleaning, size sorting, and colour sorting. Typical grade options include 7-8 mm and 8-9 mm, with Desi varieties available on request.')
add_heading(doc,'Lentils',2)
add_body(doc,'The range includes red split lentils, whole green lentils, and polished red lentils. Products are screen-cleaned and may be colour-sorted to support consistent appearance and processing performance.')
add_heading(doc,'Beans',2)
add_body(doc,'Available beans include light speckled kidney beans, pinto beans, and white/navy beans for retail, foodservice, and industrial applications.')
add_heading(doc,'Peas',2)
add_body(doc,'Whole and split green and yellow peas are supplied with controlled moisture and size grading for packing and industrial use.')

add_heading(doc,'Quality Assurance and Traceability')
add_body(doc,'Quality control is integrated throughout procurement and export preparation. Each lot is prepared against the buyer’s agreed specification, with attention to grading, cleanliness, moisture, packaging, and loading controls.')
for item in ['Sorting and grading: colour sorting, sieving, calibration, and defect reduction to maintain consistent sizing.', 'Moisture and cleanliness: target moisture levels are set by crop and buyer specification; foreign matter and broken kernels are controlled within agreed limits.', 'Sampling and traceability: lot codes and retained samples support follow-up through delivery confirmation.', 'Loading controls: stuffing photos, weight tickets, and seal information can be shared during loading.', 'Independent verification: third-party inspection and cargo insurance can be arranged when required.']:
    add_bullet(doc,item)

add_heading(doc,'Packaging and Documentation')
add_body(doc,'Packaging is selected according to product, destination, customer brand requirements, and handling conditions. Standard formats include 25 kg and 50 kg PP bags, 1 MT jumbo bags, and buyer-branded packaging where agreed. Palletisation and fumigation may be provided on request.')
add_body(doc,'A typical export document pack includes a commercial invoice, packing list, certificate of origin, and, when required, phytosanitary and fumigation certificates. Draft documents can be shared before final release for buyer review.')

add_heading(doc,'Logistics and Export Execution')
add_body(doc,'Asperlos Gostar Fardad Trading coordinates shipments from Iran to destination ports. Support may include container booking, export paperwork, loading updates, port handling, and document release.')
add_label_value(doc,[
    ('Incoterms', 'EXW, FOB, CFR and CIF, subject to quotation and destination'),
    ('Shipment modes', '20 ft and 40 ft dry containers; bulk or liner arrangements where available'),
    ('Typical lead time', 'Approximately 2-4 weeks after order confirmation, subject to crop availability and vessel schedules'),
    ('Minimum order', 'Typically one 20 ft container per SKU; mixed loads may be possible when specifications align'),
])
add_heading(doc,'Export Process',2)
for item in ['Specification and quotation: buyer shares product requirements, destination port, quantity, and incoterm; quotation is prepared accordingly.', 'Contract and payment: proforma invoice or purchase order is confirmed, payment terms are agreed, and shipment planning begins.', 'Quality preparation: product is cleaned, graded, checked, packed, and prepared against the agreed specification.', 'Stuffing and documents: loading is completed with relevant checks; draft documents are prepared for approval where applicable.', 'Shipment and follow-up: bills of lading and final documents are released as agreed, with post-arrival support when needed.']:
    add_bullet(doc,item)

add_heading(doc,'Industries and Customers Served')
for item in ['Importers and commodity traders seeking containerised or bulk purchases.', 'Retail packers and private-label buyers requiring consistent grades and packaging options.', 'Canneries, hummus producers, and food manufacturers needing ingredients suitable for industrial cooking and processing.', 'Foodservice distributors seeking dependable replenishment schedules and mixed-load options.']:
    add_bullet(doc,item)

add_heading(doc,'Commercial Approach')
add_body(doc,'We work with buyers to define the grade, physical specifications, packaging, destination port, incoterm, and delivery schedule before shipment. Commercial terms are agreed case by case. Product availability, detailed tolerances, pricing, and final shipment timing are confirmed in the quotation and contract for each order.')
add_heading(doc,'Company Positioning')
add_body(doc,'Asperlos Gostar Fardad Trading is positioned as a practical export partner for buyers seeking Iranian pulses with clear specifications, reliable quality controls, transparent documentation, and coordinated logistics. Our role is to simplify the movement of legumes from Iranian origin to the buyer’s destination market.')

doc.add_page_break()

# Persian profile
p=doc.add_paragraph(); rtl(p)
r=p.add_run('فارسی'); set_font(r,size=10,color=GOLD,bold=True)
add_heading(doc,'معرفی شرکت',rtl_text=True)
add_body(doc,'بازرگانی اسپرلوس گستر فرداد یک شرکت بازرگانی B2B مستقر در ایران است که در تأمین، درجه‌بندی، بسته‌بندی و صادرات حبوبات و محصولات بقولات فعالیت می‌کند. این شرکت، تولیدکنندگان و واحدهای فرآوری ایرانی را به واردکنندگان، عمده‌فروشان، بسته‌بندی‌کنندگان خرده‌فروشی و تولیدکنندگان صنایع غذایی در خاورمیانه، آسیا و آفریقا متصل می‌سازد.',True)
add_body(doc,'هدف ما تسهیل واردات حبوبات ایرانی برای خریداران بین‌المللی از طریق کیفیت پایدار، اسناد شفاف، هماهنگی عملیاتی حمل‌ونقل و ارتباط قابل اتکا از زمان استعلام تا تحویل کالا است.',True)
add_label_value(doc,[
    ('نام شرکت','بازرگانی اسپرلوس گستر فرداد'),
    ('حوزه فعالیت','کالاهای کشاورزی و مواد اولیه صنایع غذایی'),
    ('فعالیت اصلی','تأمین و صادرات حبوبات و محصولات بقولات'),
    ('مقر فعالیت','ایران'),
    ('بازارهای هدف','خاورمیانه، آسیا و آفریقا'),
    ('مشتریان','واردکنندگان، تجار، عمده‌فروشان، بسته‌بندی‌کنندگان، کنسروسازان، تولیدکنندگان مواد غذایی و توزیع‌کنندگان فودسرویس'),
],True)

add_heading(doc,'ماموریت، چشم‌انداز و ارزش‌ها',rtl_text=True)
add_heading(doc,'ماموریت',2,True)
add_body(doc,'تسهیل واردات حبوبات ایرانی از طریق کیفیت یکنواخت، اسناد شفاف و برنامه‌های ارسال قابل اعتماد.',True)
add_heading(doc,'چشم‌انداز',2,True)
add_body(doc,'تبدیل شدن به شریک تجاری بلندمدت و مورد اعتماد خریدارانی که به تأمین مطمئن حبوبات ایرانی، مشخصات روشن و اجرای حرفه‌ای صادرات نیاز دارند.',True)
add_heading(doc,'ارزش‌ها',2,True)
for item in ['کیفیت در اولویت هر محموله قرار دارد.', 'ارسال به‌موقع همراه با اطلاع‌رسانی مرحله‌ای.', 'شفافیت در اسناد و پشتیبانی از الزامات انطباق.', 'تمرکز بر همکاری‌های بلندمدت به جای معاملات مقطعی.', 'پاسخ‌گویی مؤثر در مناطق زمانی مختلف.']:
    add_bullet(doc,item,True)

add_heading(doc,'محصولات اصلی',rtl_text=True)
add_body(doc,'محصولات با توجه به مشخصات مورد توافق، نوع بسته‌بندی و نیاز بازار مقصد برای صادرات آماده می‌شوند.',True)
add_heading(doc,'نخود',2,True)
add_body(doc,'نخود کابلی با پاک‌سازی ماشینی، سایزبندی و سورت رنگ عرضه می‌شود. گریدهای رایج شامل ۷ تا ۸ میلی‌متر و ۸ تا ۹ میلی‌متر است و نخود دسی نیز بنا به درخواست قابل تأمین است.',True)
add_heading(doc,'عدس',2,True)
add_body(doc,'سبد عدس شامل عدس قرمز شکسته، عدس سبز کامل و عدس قرمز پولیش‌شده است. محصولات با پاک‌سازی و در صورت نیاز با سورت رنگ آماده می‌شوند تا ظاهر و عملکرد فرآوری یکنواخت‌تری داشته باشند.',True)
add_heading(doc,'لوبیا',2,True)
add_body(doc,'لوبیا چیتی روشن، لوبیا پینتو و لوبیا سفید/ناوی برای بازار خرده‌فروشی، فودسرویس و کاربردهای صنعتی قابل عرضه هستند.',True)
add_heading(doc,'نخودفرنگی',2,True)
add_body(doc,'نخودفرنگی سبز و زرد، به‌صورت کامل و لپه‌ای، با کنترل رطوبت و درجه‌بندی اندازه برای بسته‌بندی و مصارف صنعتی عرضه می‌شود.',True)

add_heading(doc,'تضمین کیفیت و ردیابی',rtl_text=True)
add_body(doc,'کنترل کیفیت از مرحله تأمین تا آماده‌سازی صادرات در فرآیند کار قرار دارد. هر محموله بر اساس مشخصات مورد توافق خریدار از نظر درجه‌بندی، پاکی، رطوبت، بسته‌بندی و کنترل‌های بارگیری آماده می‌شود.',True)
for item in ['سورت و درجه‌بندی: سورت رنگ، الک، کالیبراسیون و کاهش عیوب برای حفظ یکنواختی اندازه.', 'رطوبت و پاکی: هدف رطوبت بر اساس محصول و مشخصات خریدار تعیین می‌شود و مواد خارجی و دانه‌های شکسته در محدوده توافق‌شده کنترل می‌گردد.', 'نمونه‌برداری و ردیابی: کدهای بچ و نگهداری نمونه تا تأیید تحویل، امکان پیگیری محموله را فراهم می‌کند.', 'کنترل بارگیری: تصاویر بارگیری، باسکول و اطلاعات پلمب در طول بارگیری قابل ارائه است.', 'تأیید مستقل: بازرسی شخص ثالث و بیمه بار در صورت نیاز قابل هماهنگی است.']:
    add_bullet(doc,item,True)

add_heading(doc,'بسته‌بندی و اسناد صادراتی',rtl_text=True)
add_body(doc,'بسته‌بندی با توجه به نوع محصول، کشور مقصد، نیاز برند مشتری و شرایط حمل انتخاب می‌شود. فرمت‌های متداول شامل کیسه‌های PP بیست‌وپنج و پنجاه کیلوگرمی، جامبوبگ یک تنی و بسته‌بندی با برند خریدار است. پالت‌بندی و ضدعفونی نیز بنا به درخواست قابل انجام است.',True)
add_body(doc,'مجموعه اسناد معمول صادراتی شامل فاکتور تجاری، لیست بسته‌بندی، گواهی مبدأ و در صورت نیاز گواهی‌های بهداشت گیاهی و ضدعفونی است. پیش‌نویس اسناد می‌تواند پیش از صدور نهایی برای بررسی خریدار ارسال شود.',True)

add_heading(doc,'لجستیک و اجرای صادرات',rtl_text=True)
add_body(doc,'بازرگانی اسپرلوس گستر فرداد هماهنگی ارسال کالا از ایران تا بندر مقصد را پشتیبانی می‌کند. این خدمات می‌تواند شامل رزرو کانتینر، اسناد صادراتی، اطلاع‌رسانی بارگیری، امور بندری و تحویل اسناد باشد.',True)
add_label_value(doc,[
    ('اینکوترمز','EXW، FOB، CFR و CIF بر اساس قیمت‌گذاری و مقصد'),
    ('روش حمل','کانتینر خشک ۲۰ و ۴۰ فوت؛ حمل فله یا استفاده از لاینر در صورت امکان'),
    ('زمان معمول آماده‌سازی','حدود ۲ تا ۴ هفته پس از تأیید سفارش؛ تابع موجودی محصول و برنامه کشتیرانی'),
    ('حداقل سفارش','معمولاً یک کانتینر ۲۰ فوت برای هر SKU؛ در صورت هم‌خوانی مشخصات، بار ترکیبی امکان‌پذیر است'),
],True)
add_heading(doc,'فرآیند صادرات',2,True)
for item in ['مشخصات و قیمت‌گذاری: خریدار مشخصات محصول، بندر مقصد، مقدار و اینکوترمز را اعلام می‌کند و قیمت متناسب ارائه می‌شود.', 'قرارداد و پرداخت: پروفرما یا سفارش خرید تأیید می‌شود، شرایط پرداخت توافق می‌گردد و برنامه حمل آغاز می‌شود.', 'آماده‌سازی و کنترل کیفیت: محصول پاک‌سازی، درجه‌بندی، کنترل، بسته‌بندی و مطابق مشخصات توافق‌شده آماده می‌شود.', 'بارگیری و اسناد: بارگیری با کنترل‌های مربوط انجام و در صورت نیاز پیش‌نویس اسناد برای تأیید ارائه می‌شود.', 'ارسال و پیگیری: بارنامه و اسناد نهایی طبق توافق آزاد می‌شود و پشتیبانی پس از رسیدن کالا در صورت نیاز ادامه دارد.']:
    add_bullet(doc,item,True)

add_heading(doc,'صنایع و مشتریان هدف',rtl_text=True)
for item in ['واردکنندگان و تجار کالا که به خرید کانتینری یا عمده نیاز دارند.', 'بسته‌بندی‌کنندگان خرده‌فروشی و خریداران برند خصوصی که به گرید یکنواخت و گزینه‌های بسته‌بندی نیاز دارند.', 'کنسروسازان، تولیدکنندگان حمص و صنایع غذایی که مواد اولیه مناسب پخت و فرآوری صنعتی می‌خواهند.', 'توزیع‌کنندگان فودسرویس که به برنامه تأمین مطمئن و بارهای ترکیبی نیاز دارند.']:
    add_bullet(doc,item,True)

add_heading(doc,'رویکرد تجاری و جایگاه شرکت',rtl_text=True)
add_body(doc,'ما پیش از ارسال، گرید، مشخصات فیزیکی، بسته‌بندی، بندر مقصد، اینکوترمز و برنامه تحویل را با خریدار مشخص می‌کنیم. شرایط تجاری برای هر سفارش به‌صورت موردی توافق می‌شود. موجودی، حدود دقیق مشخصات، قیمت و زمان نهایی ارسال در پیش‌فاکتور و قرارداد هر سفارش تأیید خواهد شد.',True)
add_body(doc,'بازرگانی اسپرلوس گستر فرداد خود را به‌عنوان یک شریک عملیاتی صادرات برای خریدارانی معرفی می‌کند که به حبوبات ایرانی با مشخصات روشن، کنترل کیفیت قابل اتکا، اسناد شفاف و لجستیک هماهنگ نیاز دارند. نقش ما ساده‌سازی مسیر انتقال حبوبات از مبدأ ایران تا بازار مقصد خریدار است.',True)

add_heading(doc,'یادداشت',rtl_text=True)
add_body(doc,'این پروفایل برای معرفی عمومی شرکت تهیه شده است. مشخصات فنی، حدود کیفی، قیمت، موجودی، شرایط پرداخت و برنامه حمل برای هر سفارش به‌صورت جداگانه در پیشنهاد تجاری و قرارداد تأیید می‌شود.',True)

doc.core_properties.title = 'Asperlos Gostar Fardad Trading - Company Profile'
doc.core_properties.subject = 'Bilingual company profile'
doc.core_properties.author = 'Asperlos Gostar Fardad Trading'
doc.save(OUT)
print('Company profile created.')
